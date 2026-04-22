from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_pdf():
    pdf_path = "Sprint2_Validation_Blog_Amine.pdf"
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontSize=24, spaceAfter=20, alignment=1, textColor=colors.HexColor("#1a7a3d"))
    heading_style = ParagraphStyle('HeadingStyle', parent=styles['Heading2'], fontSize=18, spaceBefore=15, spaceAfter=10, textColor=colors.HexColor("#1a7a3d"))
    
    story.append(Paragraph("Validation Sprint 2 : Gestion du Blog", title_style))
    story.append(Paragraph("Auteur : Amine Felly", styles['Heading3']))
    story.append(Spacer(1, 0.5 * inch))

    # Backlog
    story.append(Paragraph("1. Sprint Backlog - Sprint 2", heading_style))
    backlog_data = [
        ["ID", "User Story", "Priorité", "Estimation", "Statut"],
        ["US-01", "Créer une publication (Titre, Contenu, Image)", "Haute", "5h", "Terminé"],
        ["US-04", "Créer un commentaire sur une publication", "Haute", "3h", "Terminé"],
        ["US-02", "Modifier/Supprimer ses propres publications", "Haute", "3h", "Terminé"],
        ["US-03", "Liker/Disliker des publications", "Haute", "3h", "Terminé"],
        ["US-06", "Modération Admin (Supprimer tout contenu)", "Haute", "2h", "Terminé"],
        ["US-08", "Notifications temps réel (Polling)", "Moyenne", "4h", "Terminé"]
    ]
    t = Table(backlog_data, colWidths=[0.6*inch, 3.5*inch, 0.8*inch, 0.8*inch, 0.8*inch])
    t.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1a7a3d")), ('TEXTCOLOR',(0,0),(-1,0),colors.whitesmoke), ('GRID',(0,0),(-1,-1),1,colors.black)]))
    story.append(t)

    # Story Tests
    story.append(Paragraph("2. Story Tests", heading_style))
    test_data = [
        ["US-01 : Créer Post", "Acceptation (Succès)", "Refus (Échec)"],
        ["Étant donné", "Amine Felly connecté\nTitre='Ma Recette'", "Amine Felly connecté\nTitre='Hi' (trop court)"],
        ["Quand", "Clic sur 'Publier'", "Clic sur 'Publier'"],
        ["Alors", "Post ajouté au feed", "Alerte affichée"]
    ]
    t2 = Table(test_data, colWidths=[1.5*inch, 2.75*inch, 2.75*inch])
    t2.setStyle(TableStyle([('GRID',(0,0),(-1,-1),1,colors.grey), ('BACKGROUND',(0,0),(-1,0),colors.lightgrey)]))
    story.append(t2)
    
    story.append(Spacer(1, 0.2*inch))
    
    test_data_com = [
        ["US-04 : Commenter", "Acceptation (Succès)", "Refus (Échec)"],
        ["Étant donné", "Amine Felly sur un post\nSaisie: 'Génial !'", "Amine Felly sur un post\nSaisie: '' (vide)"],
        ["Quand", "Clic sur 'Envoyer'", "Clic sur 'Envoyer'"],
        ["Alors", "Commentaire ajouté", "Alerte affichée"]
    ]
    t3 = Table(test_data_com, colWidths=[1.5*inch, 2.75*inch, 2.75*inch])
    t3.setStyle(TableStyle([('GRID',(0,0),(-1,-1),1,colors.grey), ('BACKGROUND',(0,0),(-1,0),colors.lightgrey)]))
    story.append(t3)

    doc.build(story)

def generate_word():
    doc = Document()
    title = doc.add_heading('Validation Sprint 2 : Gestion du Blog', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Auteur : Amine Felly', style='Subtitle')
    
    doc.add_heading('1. Sprint Backlog', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = "ID", "User Story", "Priorité", "Estimation", "Statut"
    
    rows = [
        ["US-01", "Créer une publication", "Haute", "5h", "Terminé"],
        ["US-04", "Créer un commentaire", "Haute", "3h", "Terminé"],
        ["US-08", "Notifications temps réel", "Moyenne", "4h", "Terminé"]
    ]
    for r in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(r): row_cells[i].text = val

    doc.add_heading('2. Story Tests', level=1)
    doc.add_heading('US-01 : Créer Post', level=2)
    t1 = doc.add_table(rows=3, cols=2); t1.style = 'Table Grid'
    t1.rows[0].cells[0].text = "Acceptation (Succès): Amine Felly publie 'Ma Recette'"
    t1.rows[0].cells[1].text = "Refus (Échec): Amine Felly publie 'Hi' (trop court)"
    
    doc.add_heading('US-04 : Créer Commentaire', level=2)
    t2 = doc.add_table(rows=3, cols=2); t2.style = 'Table Grid'
    t2.rows[0].cells[0].text = "Acceptation (Succès): Amine Felly commente 'Génial !'"
    t2.rows[0].cells[1].text = "Refus (Échec): Amine Felly commente '' (vide)"

    doc.save("Validation_Sprint2_Blog_Amine.docx")

if __name__ == "__main__":
    generate_pdf()
    generate_word()
    print("Documents mis à jour pour Amine Felly.")
